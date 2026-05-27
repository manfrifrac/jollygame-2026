from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_product_brochure(doc, product_name, type_robot, highlights, specs):
    # PAGINA 1: Commerciale
    doc.add_heading(f'Bozza Depliant: {product_name} - Pagina 1 (Commerciale)', level=1)
    
    # Placeholder Immagine Hero
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ PROMPT IMMAGINE PER AI: Immagine fotorealistica ad alta risoluzione del robot da piscina '{product_name}' posizionato a bordo piscina con acqua cristallina sullo sfondo. Stile pulito, luminoso e professionale. ]")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True
    
    doc.add_heading('TITOLO E SOTTOTITOLO', level=2)
    doc.add_paragraph(f"Titolo Principale: {product_name} - L'eccellenza della pulizia {type_robot}", style='Title')
    doc.add_paragraph(f"Sottotitolo: Prestazioni superiori e affidabilità per una piscina sempre perfetta e pronta all'uso.")
    
    doc.add_heading('INTRODUZIONE', level=2)
    doc.add_paragraph("Scopri la libertà di una piscina sempre impeccabile. Progettato per garantire la massima efficienza, questo robot pulitore combina tecnologia avanzata e semplicità d'uso, permettendoti di goderti solo il meglio della tua piscina senza pensieri.")
    
    doc.add_heading('PUNTI DI FORZA', level=2)
    for hl in highlights:
        doc.add_paragraph(hl, style='List Bullet')
        
    doc.add_page_break()

    # PAGINA 2: Tecnico
    doc.add_heading(f'Bozza Depliant: {product_name} - Pagina 2 (Tecnica)', level=1)
    
    # Placeholder Immagine Tecnica
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ PROMPT IMMAGINE PER AI: Vista esplosa o infografica tecnica pulita del robot, sfondo bianco, indicatori grafici blu e grigi che mostrano il flusso d'acqua e i componenti principali. Stile vettoriale o render 3D tecnico. ]")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True

    doc.add_heading('SPECIFICHE TECNICHE', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Caratteristica'
    hdr_cells[1].text = 'Valore'
    
    for key, value in specs.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value
        
    doc.add_heading('CONTENUTO DELLA CONFEZIONE', level=2)
    doc.add_paragraph('1. Robot Pulitore', style='List Number')
    doc.add_paragraph('2. Cestello filtro a maglia fine', style='List Number')
    doc.add_paragraph('3. Unità di controllo o Caricabatterie', style='List Number')
    doc.add_paragraph('4. Manuale d\'uso e garanzia', style='List Number')

    doc.add_heading('DISTRIBUZIONE', level=3)
    p = doc.add_paragraph("Distribuito da JollyGame. Tutte le specifiche sono soggette a modifiche per miglioramenti tecnici.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

doc = Document()

# ZODIAC MX8 (Idraulico)
add_product_brochure(doc, 
    "Zodiac MX8 BARACUDA", "idraulica",
    ["Sistema di Navigazione X-Drive: Copertura totale della piscina.", "Aspirazione Ciclonica: Motore potente per detriti ostinati.", "Cingoli di Trazione: Perfetta aderenza su tutti i rivestimenti."],
    {"Tipologia": "Idraulico ad aspirazione", "Zone di pulizia": "Fondo e pareti", "Collegamento": "Skimmer o presa aspirafango", "Potenza pompa richiesta": "Minimo 3/4 CV"}
)

# NIYA TRACKER 35 (Cordless ipotetico)
add_product_brochure(doc, 
    "Niya Tracker 35", "senza fili",
    ["Libertà Cordless: Nessun cavo, nessuna restrizione.", "Sensori Intelligenti: Riconosce le pareti e inverte la marcia.", "Batteria Lunga Durata: Fino a 90 minuti di autonomia."],
    {"Tipologia": "Batteria ricaricabile", "Autonomia": "Fino a 90 minuti", "Filtrazione": "Cestello estraibile dall'alto", "Peso": "5.5 kg"}
)

# GRE ELECTRIC VAC
add_product_brochure(doc, 
    "Gre Electric Vac", "pratica e veloce",
    ["Plug & Play: Subito pronto all'uso.", "Design Compatto: Facile da maneggiare e svuotare.", "Filtrazione Efficiente: Trattiene anche le particelle fini."],
    {"Tipologia": "Elettrico standard", "Fondo": "Fondo piatto o lieve pendenza", "Ciclo di pulizia": "1.5 ore / 2 ore", "Rivestimento": "Liner, PVC, Cemento"}
)

# GRE WERUNNER PLUS
add_product_brochure(doc, 
    "Gre WERUNNER PLUS", "avanzata",
    ["Massima Copertura: Ideale per piscine di medie e grandi dimensioni.", "Spazzolatura Attiva: Rimuove alghe e batteri dal fondo.", "Accesso Rapido al Filtro: Pulizia senza sporcarsi le mani."],
    {"Tipologia": "Elettrico automatico", "Zone di pulizia": "Fondo", "Cavo": "15 metri", "Filtrazione": "Filtro a cartuccia lavabile"}
)

doc.save('Bozze_Depliant_Mancanti.docx')
print("Document created successfully.")
