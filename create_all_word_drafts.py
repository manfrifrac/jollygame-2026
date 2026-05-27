from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_product_brochure(doc, product_name, type_robot, highlights, specs):
    doc.add_heading(f'Bozza Depliant: {product_name} - Pagina 1 (Commerciale)', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ PROMPT IMMAGINE PER AI: Immagine fotorealistica ad alta risoluzione del robot da piscina '{product_name}' posizionato a bordo piscina con acqua cristallina sullo sfondo. Stile pulito, luminoso e professionale. ]")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True
    
    doc.add_heading('TITOLO E SOTTOTITOLO', level=2)
    doc.add_paragraph(f"Titolo Principale: {product_name} - L'eccellenza della pulizia {type_robot}", style='Title')
    doc.add_paragraph(f"Sottotitolo: Prestazioni superiori e affidabilità per una piscina sempre perfetta e pronta all'uso.")
    
    doc.add_heading('INTRODUZIONE', level=2)
    doc.add_paragraph("Scopri la libertà di una piscina sempre impeccabile. Progettato per garantire la massima efficienza, questo robot pulitore combina tecnologia avanzata e semplicità d'uso, permettendoti di goderti solo il meglio della tua piscina senza pensieri.")
    
    doc.add_heading('PUNTI DI FORZA', level=2)
    for hl in highlights:
        doc.add_paragraph(hl, style='List Bullet')
        
    doc.add_page_break()

    doc.add_heading(f'Bozza Depliant: {product_name} - Pagina 2 (Tecnica)', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ PROMPT IMMAGINE PER AI: Vista esplosa o infografica tecnica pulita del robot, sfondo bianco, indicatori grafici blu e grigi che mostrano il flusso d'acqua e i componenti principali. Stile vettoriale o render 3D tecnico. ]")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True

    doc.add_heading('SPECIFICHE TECNICHE', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Caratteristica'
    hdr_cells[1].text = 'Valore'
    
    for key, value in specs.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value
        
    doc.add_heading('CONTENUTO DELLA CONFEZIONE', level=2)
    doc.add_paragraph('1. Robot Pulitore', style='List Number')
    doc.add_paragraph('2. Cestello filtro a maglia fine', style='List Number')
    doc.add_paragraph('3. Unità di controllo o Caricabatterie', style='List Number')
    doc.add_paragraph('4. Manuale d\'uso e garanzia', style='List Number')

    doc.add_heading('DISTRIBUZIONE', level=3)
    p = doc.add_paragraph("Distribuito da JollyGame. Tutte le specifiche sono soggette a modifiche per miglioramenti tecnici.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

doc = Document()

all_products = [
    ("ZODIAC FREERIDER RF 5200 iQ", "cordless", ["Libertà Senza Fili", "Aspirazione Ciclonica", "Agilità su ogni rivestimento"], {"Tipo": "Batteria", "Zona": "Fondo/Pareti/Linea", "Ciclo": "2.5h"}),
    ("ZODIAC VOYAGER RE 4400 iQ", "intelligente", ["Sensore Nav System", "Aspirazione Ciclonica", "App iAqualink"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea", "Cavo": "18m"}),
    ("ZODIAC VOYAGER RE 45AM iQ", "intelligente", ["Filtrazione 4L", "Agilità di movimento", "Design Moderno"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea", "Cavo": "18m"}),
    ("ZODIAC VOYAGER RE 4300", "potente", ["Filtrazione Progressiva", "Spazzolatura Attiva", "Accesso al filtro"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea", "Cavo": "15m"}),
    ("ZODIAC OT 3200 Tornax", "compatta", ["Ultra leggero", "Facile manutenzione", "Navigazione intelligente"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti", "Peso": "5.5kg"}),
    ("ZODIAC ALPHA iQ RA 6700 iQ", "premium", ["Sensor Nav System", "Aspirazione Ultra Potente", "App Control Avanzato"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea", "Sensor": "Sì"}),
    ("ZODIAC VOYAGER RE 4700 iQ", "avanzata", ["Filtrazione Doppia", "Indicatori LED", "Controllo App"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea", "Filtro": "Doppio"}),
    ("ZODIAC ALPHA iQ RA 6570 iQ", "premium", ["Prestazioni Massime", "Pulizia Personalizzata", "Sistema Lift System"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea", "App": "Sì"}),
    ("ZODIAC Tornax PRO RT 3200", "professionale", ["Affidabilità Pro", "Spazzole dedicate", "Pistola di aspirazione"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti", "Serie": "PRO"}),
    ("ZODIAC XA TYPE XA 2010", "agile", ["Navigazione Smart", "Filtrazione efficiente", "Leggerezza"], {"Tipo": "Elettrico", "Zona": "Fondo", "Serie": "XA"}),
    ("ZODIAC TornaX PRO RT 2100", "compatta", ["Prezzo imbattibile", "Qualità Zodiac", "Manutenzione facile"], {"Tipo": "Elettrico", "Zona": "Fondo", "Peso": "5.5kg"}),
    ("ZODIAC MX8 BARACUDA", "idraulica", ["X-Drive Navigation", "Turbina Potente", "Cingoli di Trazione"], {"Tipo": "Idraulico", "Pompa": "Min 3/4 CV", "Zona": "Fondo/Pareti"}),
    ("Gre WERUNNER PLUS", "automatica", ["Copertura Totale", "Spazzole PVC", "Filtro a cartuccia"], {"Tipo": "Elettrico", "Zona": "Fondo", "Cavo": "15m"}),
    ("Gre Electric Vac", "manuale", ["Aspirazione immediata", "Nessun cavo", "Batteria ricaricabile"], {"Tipo": "Manuale Batteria", "Uso": "SPA/Piccole piscine", "Batteria": "Litio"}),
    ("Gre action vac", "manuale", ["Massima manovrabilità", "Filtro integrato", "Uso facile"], {"Tipo": "Manuale", "Fondo": "Tutti i tipi", "Peso": "2kg"}),
    ("niya tracker 25", "cordless", ["Total Cordless", "Smart Scan", "Filtrazione 3.5L"], {"Tipo": "Batteria", "Autonomia": "60 min", "Zona": "Fondo"}),
    ("niya tracker 35", "cordless", ["Batteria Lunga Durata", "Recupero Facile", "Potenza aspirazione"], {"Tipo": "Batteria", "Autonomia": "90 min", "Zona": "Fondo"}),
    ("Niya Sonar 30", "sonica", ["Navigazione a ultrasuoni", "Copertura efficiente", "Cordless"], {"Tipo": "Batteria", "Spec": "Sensore Sonar", "Zona": "Fondo"}),
    ("Niya Eclipse 35", "oscura", ["Design elegante", "Pulizia profonda", "Senza fili"], {"Tipo": "Batteria", "Design": "Eclissi", "Zona": "Fondo"})
]

for name, robot_type, highlights, specs in all_products:
    add_product_brochure(doc, name, robot_type, highlights, specs)

doc.save('Tutti_i_Depliant_JollyGame.docx')
print("Complete document created successfully.")
