import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Elenco file IT esistenti per controllo rapido
existing_it_files = [
    "OT_3200_Depliant_OT_3200_IT.pdf",
    "RA_6700_iQ_Depliant_RA_6700_iQ_IT.pdf",
    "RE_4300_Depliant_RE_4300_IT.pdf",
    "RE_4400_iQ_Depliant_RE_4400_iQ_IT.pdf",
    "RE_4700_iQ_Depliant_RE_4700_iQ_IT.pdf",
    "RF_5200_iQ_Leaflet_RF_5200_iQ_IT.pdf"
]

def has_existing_brochure(product_name):
    # Mapping semplice
    mapping = {
        "RF 5200 iQ": "RF_5200_iQ_Leaflet_RF_5200_iQ_IT.pdf",
        "RE 4400 iQ": "RE_4400_iQ_Depliant_RE_4400_iQ_IT.pdf",
        "RE 4300": "RE_4300_Depliant_RE_4300_IT.pdf",
        "OT 3200": "OT_3200_Depliant_OT_3200_IT.pdf",
        "RA 6700 iQ": "RA_6700_iQ_Depliant_RA_6700_iQ_IT.pdf",
        "RE 4700 iQ": "RE_4700_iQ_Depliant_RE_4700_iQ_IT.pdf"
    }
    for key, filename in mapping.items():
        if key.lower() in product_name.lower():
            return filename
    return None

def add_full_brochure(doc, product_name, type_robot, highlights, specs):
    doc.add_heading(f'Bozza Depliant: {product_name} - Pagina 1 (Commerciale)', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ PROMPT IMMAGINE PER AI: Immagine fotorealistica ad alta risoluzione del robot da piscina '{product_name}' posizionato a bordo piscina con acqua cristallina sullo sfondo. ]")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True
    doc.add_heading('TITOLO E SOTTOTITOLO', level=2)
    doc.add_paragraph(f"Titolo Principale: {product_name} - Pulizia {type_robot}", style='Title')
    doc.add_heading('PUNTI DI FORZA', level=2)
    for hl in highlights:
        doc.add_paragraph(hl, style='List Bullet')
    doc.add_page_break()
    doc.add_heading(f'Bozza Depliant: {product_name} - Pagina 2 (Tecnica)', level=1)
    doc.add_heading('SPECIFICHE TECNICHE', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for key, value in specs.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value
    doc.add_page_break()

doc = Document()
doc.add_heading('Catalogo Depliant JollyGame 2026', level=0)

all_products = [
    ("ZODIAC FREERIDER RF 5200 iQ", "cordless", ["Senza Fili", "Ciclonico"], {"Tipo": "Batteria", "Zona": "Fondo/Pareti/Linea"}),
    ("ZODIAC VOYAGER RE 4400 iQ", "intelligente", ["Sensor Nav System", "iAqualink"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea"}),
    ("ZODIAC VOYAGER RE 45AM iQ", "intelligente", ["Design Moderno", "Potente"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea"}),
    ("ZODIAC VOYAGER RE 4300", "potente", ["Filtrazione Progressiva"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea"}),
    ("ZODIAC OT 3200 Tornax", "compatta", ["Ultra leggero"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti"}),
    ("ZODIAC ALPHA iQ RA 6700 iQ", "premium", ["Top di gamma"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea"}),
    ("ZODIAC VOYAGER RE 4700 iQ", "avanzata", ["Filtrazione Doppia"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea"}),
    ("ZODIAC ALPHA iQ RA 6570 iQ", "premium", ["Personalizzabile"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti/Linea"}),
    ("ZODIAC Tornax PRO RT 3200", "professionale", ["Affidabilità"], {"Tipo": "Elettrico", "Zona": "Fondo/Pareti"}),
    ("ZODIAC XA TYPE XA 2010", "agile", ["Smart"], {"Tipo": "Elettrico", "Zona": "Fondo"}),
    ("ZODIAC TornaX PRO RT 2100", "compatta", ["Essenziale"], {"Tipo": "Elettrico", "Zona": "Fondo"}),
    ("ZODIAC MX8 BARACUDA", "idraulica", ["X-Drive", "Cingoli"], {"Tipo": "Idraulico", "Pompa": "Min 3/4 CV"}),
    ("Gre WERUNNER PLUS", "automatica", ["Efficace"], {"Tipo": "Elettrico", "Zona": "Fondo"}),
    ("Gre Electric Vac", "manuale", ["Cordless"], {"Tipo": "Batteria", "Uso": "SPA"}),
    ("Gre action vac", "manuale", ["Leggero"], {"Tipo": "Manuale", "Peso": "2kg"}),
    ("niya tracker 25", "cordless", ["Senza Fili"], {"Tipo": "Batteria", "Zona": "Fondo"}),
    ("niya tracker 35", "cordless", ["Lunga durata"], {"Tipo": "Batteria", "Zona": "Fondo"}),
    ("Niya Sonar 30", "sonica", ["Sonar"], {"Tipo": "Batteria", "Zona": "Fondo"}),
    ("Niya Eclipse 35", "cordless", ["Elegante"], {"Tipo": "Batteria", "Zona": "Fondo"})
]

for name, robot_type, highlights, specs in all_products:
    existing_file = has_existing_brochure(name)
    if existing_file:
        doc.add_heading(f'PRODOTTO: {name}', level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"DEPLIANT ESISTENTE DISPONIBILE: {existing_file}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0) # Verde
        doc.add_paragraph(f"Nota: Per questo modello utilizzare il PDF ufficiale presente nella cartella 'depliant_it_short'. Non è necessaria la generazione AI.")
        doc.add_paragraph("-" * 50)
    else:
        add_full_brochure(doc, name, robot_type, highlights, specs)

doc.save('Depliant_JollyGame_Finali.docx')
print("Refined document created successfully.")
